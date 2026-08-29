"""Shared data-layer logic used by the JSON API (api.py)."""
import calendar
import datetime
import json
import re
import uuid
import zipfile
from io import BytesIO
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import bleach
from bleach.css_sanitizer import CSSSanitizer
from flask_login import current_user

from sqlalchemy.dialects.postgresql import insert as pg_insert

from extensions import db
from models import User, UserSettings, Folder, Chapter, ChapterVersion, ChapterWritingActivity, Goal, GoalType, GoalPeriodHistory
from permissions import accessible_book_ids

VERSION = "0.21.1"

CHAPTER_VERSION_MIN_INTERVAL = datetime.timedelta(minutes=5)
CHAPTER_VERSION_RETENTION = 50

STATS_STALE_GAP = datetime.timedelta(days=7)

# Caps how much of a single heartbeat gap gets credited as "active" writing
# time -- 3x the 20s client heartbeat interval, enough to absorb one missed
# beat without letting a laptop-sleep/backgrounded-tab gap (which can be
# hours) get counted as active time once the next heartbeat lands.
MAX_HEARTBEAT_ELAPSED = datetime.timedelta(seconds=60)

# Days without writing activity before an incomplete chapter is surfaced as
# "stale" on its folder's stats page.
STALE_CHAPTER_DAYS = 14

# P1.2 Book Types -- metadata (Folder.book_type) layered on the existing
# Book/Folder/Chapter structure, never a separate model per type.
BOOK_TYPES = ('general', 'novel', 'journal', 'documentation')

# WPM from a handful of words or a few seconds of activity is noise, not a
# meaningful rate -- both thresholds must be met before a WPM figure is
# calculated at all (see calculate_wpm). Centralized here, the one logical
# home for both the single-user and per-contributor WPM call sites, rather
# than letting each surface re-derive its own cutoff.
MIN_WPM_TYPED_WORDS = 25
MIN_WPM_ACTIVE_SECONDS = 60


def calculate_wpm(words_typed: int, active_seconds: int) -> float | None:
    """The one WPM calculation every surface (single-user and per-contributor
    alike) shares. None below MIN_WPM_TYPED_WORDS/MIN_WPM_ACTIVE_SECONDS --
    both thresholds must be met -- rather than a misleadingly precise figure
    from a tiny sample. Deliberately keeps using gross `words_typed`, never
    net of words_deleted: WPM measures genuine typing activity, not net
    manuscript growth, and pasted words are never included in words_typed to
    begin with (see ChapterWritingActivity's docstring)."""
    if words_typed < MIN_WPM_TYPED_WORDS or active_seconds < MIN_WPM_ACTIVE_SECONDS:
        return None
    return round(words_typed / (active_seconds / 60), 1)


def words_written_from(words_typed: int, words_deleted: int) -> int:
    """The "Words written" stat shown everywhere in the UI: net genuine
    typing after genuine deletion, never negative. Derived on read, not a
    stored column -- see ChapterWritingActivity.words_deleted's docstring."""
    return max(words_typed - words_deleted, 0)


def timezone_info(timezone_name: str | None) -> ZoneInfo:
    """Resolve a stored IANA zone, falling back to UTC only for legacy users
    whose browser has not recorded one yet. Invalid non-null values are never
    persisted by the API; the defensive fallback keeps a manually corrupted
    row from breaking every calendar-aware endpoint."""
    try:
        return ZoneInfo(timezone_name or 'UTC')
    except (ZoneInfoNotFoundError, ValueError):
        return ZoneInfo('UTC')


def is_valid_timezone(timezone_name) -> bool:
    if not isinstance(timezone_name, str) or not timezone_name or len(timezone_name) > 64:
        return False
    try:
        ZoneInfo(timezone_name)
        return True
    except (ZoneInfoNotFoundError, ValueError):
        return False


def user_local_datetime(user=None, at_utc: datetime.datetime | None = None) -> datetime.datetime:
    """An absolute UTC instant interpreted in `user`'s IANA timezone."""
    user = user or current_user
    instant = at_utc or datetime.datetime.now(datetime.timezone.utc)
    if instant.tzinfo is None:
        instant = instant.replace(tzinfo=datetime.timezone.utc)
    return instant.astimezone(timezone_info(getattr(user, 'timezone', None)))


def user_local_today(user=None, at_utc: datetime.datetime | None = None) -> datetime.date:
    return user_local_datetime(user, at_utc).date()


def local_date_bounds_utc(
    start_date: datetime.date,
    end_date: datetime.date,
    user,
) -> tuple[datetime.datetime, datetime.datetime]:
    """Inclusive local calendar-date range expressed as a half-open UTC
    interval. Constructing the next local midnight through ZoneInfo handles
    23/25-hour DST days without any custom offset arithmetic."""
    zone = timezone_info(getattr(user, 'timezone', None))
    start = datetime.datetime.combine(start_date, datetime.time.min, tzinfo=zone)
    after_end = datetime.datetime.combine(end_date + datetime.timedelta(days=1), datetime.time.min, tzinfo=zone)
    return start.astimezone(datetime.timezone.utc), after_end.astimezone(datetime.timezone.utc)


def prune_words_per_day(words_per_day: dict, today: datetime.date | None = None) -> dict:
    """Drop zero-word days, which just mean nothing was last-touched that
    day rather than "0 words were written" -- except today, which stays
    even at 0 so the chart doesn't just stop looking alive. If today itself
    is 0 and it's been more than a week since the last day with any real
    words (or there's never been one), drop today too, rather than show a
    lone empty bar stranded past a long gap of inactivity."""
    real_days = sorted(d for d, c in words_per_day.items() if c > 0)
    pruned = {d: words_per_day[d] for d in real_days}

    today = today or user_local_today()
    today_iso = today.isoformat()
    if today_iso not in words_per_day or words_per_day[today_iso] > 0:
        return pruned
    if real_days and today - datetime.date.fromisoformat(real_days[-1]) <= STATS_STALE_GAP:
        pruned[today_iso] = words_per_day[today_iso]
    return pruned


def clean_name(name: str) -> str:
    """Light cosmetic cleanup for a user-supplied name (no filesystem safety needed --
    names are DB values matched by exact string, never used to build a path)."""
    return name.strip()[:255]


def _to_alpha(n: int) -> str:
    """1 -> a, 2 -> b, ..., 26 -> z, 27 -> aa, ... (bijective base-26)."""
    s = ""
    while n > 0:
        n, r = divmod(n - 1, 26)
        s = chr(97 + r) + s
    return s


def _to_roman(n: int) -> str:
    vals = [
        (1000, "m"), (900, "cm"), (500, "d"), (400, "cd"), (100, "c"), (90, "xc"),
        (50, "l"), (40, "xl"), (10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i"),
    ]
    s = ""
    for value, symbol in vals:
        count, n = divmod(n, value)
        s += symbol * count
    return s


_ORDERED_MARKERS = [lambda n: f"{n}.", lambda n: f"{_to_alpha(n)}.", lambda n: f"{_to_roman(n)}."]
_UNORDERED_MARKERS = ["•", "◦", "▪"]  # •, ◦, ▪


def _list_marker(depth: int, ordered: bool, index: int, checked) -> str:
    """The prefix for one list item -- cycles decimal/alpha/roman (ordered)
    or disc/circle/square (unordered) every 3 nesting levels, matching the
    editor's CSS, or a checkbox glyph for a checklist item regardless of
    ordered/unordered (checklists are always rendered as a plain ul)."""
    if checked is not None:
        return "[x]" if checked else "[ ]"
    if ordered:
        return _ORDERED_MARKERS[depth % 3](index)
    return _UNORDERED_MARKERS[depth % 3]


def _walk_list(list_elem, depth=0):
    """Yield (depth, ordered, index, checked, own_children) for every <li>
    under list_elem, in document order, recursing into a nested <ul>/<ol>
    right after its parent <li>'s own content (own_children excludes any
    nested list, which is walked separately by the recursive call). index is
    1-based and resets within each nested list. Each consumer picks its own
    marker convention from these fields -- e.g. depth-cycling glyphs for
    docx/rtf/txt (mirroring the editor's own display), vs. markdown's flat
    "-"/"N." convention where indentation alone conveys nesting."""
    ordered = list_elem.name == "ol"
    index = 0
    for li in list_elem.find_all("li", recursive=False):
        index += 1
        classes = li.get("class") or []
        checked = ("checked" in classes) if "checklist-item" in classes else None
        own_children = [c for c in li.children if getattr(c, "name", None) not in ("ul", "ol")]
        yield depth, ordered, index, checked, own_children
        for nested in li.find_all(["ul", "ol"], recursive=False):
            yield from _walk_list(nested, depth + 1)


def html_to_text(html: str) -> str:
    """Convert HTML to plain text, with list items prefixed by their marker
    (bullet/number/checkbox) and indented by nesting depth -- otherwise a
    list is indistinguishable from a run of separate paragraphs."""
    soup = BeautifulSoup(html, "html.parser")

    def block_text(elem) -> str:
        if elem.name in ("ul", "ol"):
            lines = [
                f"{'  ' * depth}{_list_marker(depth, ordered, index, checked)} "
                f"{''.join(c.get_text() for c in own).strip()}"
                for depth, ordered, index, checked, own in _walk_list(elem)
            ]
            return "\n".join(lines)
        return elem.get_text(separator="\n")

    parts = []
    for elem in soup.children:
        text = (elem if isinstance(elem, str) else block_text(elem)).strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


# Search 2.0 (P1.1) -----------------------------------------------------
#
# html_to_search_text is deliberately a *different* flattening than
# html_to_text above: html_to_text injects bullet/number/checkbox markers
# for export/WPM word-counting, which are display conventions a user isn't
# literally searching for and would otherwise make an occurrence's offset
# meaningless relative to what's actually on screen. Block-level tags (and
# <br>) become a plain "\n" boundary here instead, and inline formatting
# tags contribute no boundary at all, so "<strong>Talak</strong>tei" stays
# one contiguous, searchable "Talaktei" -- see the frontend's matching
# canonicalSearchText (utils/searchText.ts), which MUST walk the live
# editor DOM with the exact same block/inline rules for a server-found
# occurrence offset to resolve to the right place client-side.
_SEARCH_BLOCK_TAGS = {"p", "div", "li", "ul", "ol", "hr"}


def html_to_search_text(html: str) -> str:
    """Canonical searchable plain-text view of Chapter content HTML for
    Search 2.0 occurrence matching. Never used for export or word counts --
    see html_to_text for that. Does not alter/re-save the stored HTML."""
    soup = BeautifulSoup(html, "html.parser")

    def walk(node) -> str:
        if isinstance(node, str):
            return str(node)
        if node.name == "br":
            return "\n"
        if node.name == "img":
            return ""
        text = "".join(walk(c) for c in node.children)
        return text + "\n" if node.name in _SEARCH_BLOCK_TAGS else text

    return "".join(walk(c) for c in soup.contents)


def find_literal_occurrences(haystack: str, needle: str) -> list[tuple[int, int]]:
    """Case-insensitive, literal (not user-facing regex) substring match
    offsets in `haystack`, left to right, non-overlapping. Matching happens
    directly against the original (not lower-cased) string via a
    re.escape'd, re.IGNORECASE pattern -- not haystack.lower().find(...) --
    specifically so Unicode case-folding that changes a character's length
    (rare, but real: e.g. 'İ'.lower() in some locales) can never desync a
    match span from the original text's actual offsets. re.escape makes any
    query safe to compile, including regex-special characters like
    `(`, `[`, `.`, `*`."""
    if not needle:
        return []
    pattern = re.compile(re.escape(needle), re.IGNORECASE)
    return [(m.start(), m.end()) for m in pattern.finditer(haystack)]


SEARCH_SNIPPET_CONTEXT_CHARS = 90


def _normalize_snippet_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def build_search_snippet(text: str, start: int, end: int, *, context: int = SEARCH_SNIPPET_CONTEXT_CHARS) -> dict:
    """Display-only context window around one match -- never touches the
    underlying text/offsets used for matching or navigation. Trims to a
    nearby word boundary when one exists within the window (a run-on word
    longer than the whole window is left as-is, per "when practical"), and
    collapses whitespace/newline runs so a snippet spanning a paragraph
    break still reads as one clean line."""
    window_start = max(0, start - context)
    window_end = min(len(text), end + context)
    before_raw = text[window_start:start]
    after_raw = text[end:window_end]

    leading_ellipsis = window_start > 0
    trailing_ellipsis = window_end < len(text)

    if leading_ellipsis:
        m = re.search(r"\s", before_raw)
        if m:
            before_raw = before_raw[m.end():]
    if trailing_ellipsis:
        ws_matches = list(re.finditer(r"\s", after_raw))
        if ws_matches:
            after_raw = after_raw[:ws_matches[-1].start()]

    return {
        "before": _normalize_snippet_whitespace(before_raw),
        "match": text[start:end],
        "after": _normalize_snippet_whitespace(after_raw),
        "leadingEllipsis": leading_ellipsis,
        "trailingEllipsis": trailing_ellipsis,
    }


def search_chapter_matches(chapter: Chapter, query: str) -> list[dict]:
    """Every literal match in one chapter's title, content, and notes, in
    title -> content (document order) -> notes (notes order), each carrying
    enough to build an API match row: source, a source-local zero-based
    occurrenceIndex (None for title, which is always at most one result),
    and the offsets plus canonical text they were found in (for snippet
    generation by the caller, which also batches Book metadata -- kept out
    of this function to avoid a query per occurrence)."""
    matches = []

    title_hits = find_literal_occurrences(chapter.name, query)
    if title_hits:
        start, end = title_hits[0]
        matches.append({"source": "title", "occurrenceIndex": None, "startOffset": start, "endOffset": end, "text": chapter.name})

    content_text = html_to_search_text(chapter.content_html)
    for i, (start, end) in enumerate(find_literal_occurrences(content_text, query)):
        matches.append({"source": "content", "occurrenceIndex": i, "startOffset": start, "endOffset": end, "text": content_text})

    notes_text = chapter.notes_text or ""
    for i, (start, end) in enumerate(find_literal_occurrences(notes_text, query)):
        matches.append({"source": "notes", "occurrenceIndex": i, "startOffset": start, "endOffset": end, "text": notes_text})

    return matches


_CSS_SANITIZER = CSSSanitizer(allowed_css_properties=["text-indent", "text-align"])


def sanitize_html(html: str) -> str:
    """Strip unwanted tags to prevent script injection. A CSSSanitizer is
    required here -- without one, bleach silently drops the *entire* style
    attribute (not just disallowed properties) even though "style" is listed
    as an allowed attribute below, which was quietly discarding the editor's
    first-line-indent and text-align formatting on every save."""
    allowed_tags = [
        "b", "strong", "i", "em", "u", "p", "br", "div", "ul", "ol", "li", "hr",
        "span", "img", "a"
    ]
    allowed_attrs = {
        "*": ["class", "style"],
        "img": ["src", "alt", "width", "height", "style"],
        # Internal-reference identity lives in these data attributes. The
        # visible link text and href are presentation/navigation aids only;
        # target type + UUID remain stable across renames and moves.
        "a": [
            "href", "title", "rel", "target",
            "data-calwriter-target-type", "data-calwriter-target-id",
        ],
    }
    return bleach.clean(
        html,
        tags=allowed_tags,
        attributes=allowed_attrs,
        protocols=["http", "https", "mailto", "calwriter"],
        css_sanitizer=_CSS_SANITIZER,
        strip=True,
    )


def append_html_to_docx(doc: Document, html: str) -> None:
    """Append limited HTML content to a DOCX document."""
    soup = BeautifulSoup(html, "html.parser")

    def process(elem, paragraph, formatting=None):
        if formatting is None:
            formatting = {}
        if isinstance(elem, str):
            run = paragraph.add_run(elem)
            run.bold = formatting.get("bold", False)
            run.italic = formatting.get("italic", False)
            run.underline = formatting.get("underline", False)
            return
        tag = elem.name
        fmt = formatting.copy()
        if tag in ("strong", "b"):
            fmt["bold"] = True
        if tag in ("em", "i"):
            fmt["italic"] = True
        if tag == "u":
            fmt["underline"] = True
        if tag == "img":
            src = elem.get("src", "")
            if src.startswith("data:image/"):
                import base64
                header, b64 = src.split(",", 1)
                data = base64.b64decode(b64)
                bio = BytesIO(data)
                width = elem.get("width")
                height = elem.get("height")
                w = Inches(int(width) / 96) if width and width.isdigit() else None
                h = Inches(int(height) / 96) if height and height.isdigit() else None
                doc.add_picture(bio, width=w, height=h)
            return
        if tag in ("p", "div", "br"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            for child in elem.children:
                process(child, p, fmt)
            return
        if tag in ("ul", "ol"):
            for depth, ordered, index, checked, own in _walk_list(elem):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 * (depth + 1))
                p.add_run(f"{_list_marker(depth, ordered, index, checked)}  ")
                for child in own:
                    process(child, p, fmt)
            return
        for child in elem.children:
            process(child, paragraph, fmt)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    for child in soup.children:
        process(child, p)


def render_chapter_docx(content_html: str) -> Document:
    doc = Document()
    append_html_to_docx(doc, content_html)
    return doc


def render_book_docx(chapters: list) -> Document:
    doc = Document()
    for idx, chapter in enumerate(chapters):
        doc.add_heading(chapter.name, level=1)
        append_html_to_docx(doc, chapter.content_html)
        if idx < len(chapters) - 1:
            doc.add_page_break()
    return doc


def html_to_markdown(html: str) -> str:
    """Convert the editor's limited HTML dialect to Markdown. Alignment/indent
    styling has no plain-Markdown equivalent and is dropped, same as .txt."""
    soup = BeautifulSoup(html, "html.parser")

    def render_inline(elem) -> str:
        if isinstance(elem, str):
            return elem
        tag = elem.name
        inner = "".join(render_inline(c) for c in elem.children)
        if tag in ("strong", "b"):
            return f"**{inner}**" if inner.strip() else inner
        if tag in ("em", "i"):
            return f"*{inner}*" if inner.strip() else inner
        if tag == "u":
            return f"<u>{inner}</u>" if inner.strip() else inner
        if tag == "br":
            return "  \n"
        return inner

    def render_children(children) -> str:
        return "".join(render_inline(c) for c in children)

    blocks = []
    for elem in soup.children:
        if isinstance(elem, str):
            text = elem.strip()
            if text:
                blocks.append(text)
            continue
        tag = elem.name
        if tag == "hr":
            blocks.append("---")
        elif tag in ("ul", "ol"):
            items = []
            for depth, ordered, index, checked, own in _walk_list(elem):
                if checked is not None:
                    prefix = "- [x]" if checked else "- [ ]"
                elif ordered:
                    prefix = f"{index}."
                else:
                    prefix = "-"
                items.append(f"{'  ' * depth}{prefix} {render_children(own).strip()}")
            blocks.append("\n".join(items))
        else:
            text = render_inline(elem).strip()
            if text:
                blocks.append(text)
    return "\n\n".join(blocks)


def _rtf_escape(text: str) -> str:
    text = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
    return "".join(ch if ord(ch) < 128 else f"\\u{ord(ch)}?" for ch in text)


def html_to_rtf_body(html: str) -> str:
    """RTF content fragment for one chapter's HTML, without the document
    header/footer -- render_chapter_rtf / render_folder_rtf wrap this."""
    soup = BeautifulSoup(html, "html.parser")

    def render(elem) -> str:
        if isinstance(elem, str):
            return _rtf_escape(elem)
        tag = elem.name
        if tag == "br":
            return "\\line "
        if tag == "hr":
            return "\\line \\emdash\\emdash\\emdash\\emdash\\emdash\\line "
        if tag in ("ul", "ol"):
            parts = []
            for depth, ordered, index, checked, own in _walk_list(elem):
                marker = _rtf_escape(_list_marker(depth, ordered, index, checked))
                indent = 360 * (depth + 1)
                inner = "".join(render(c) for c in own)
                parts.append(f"\\li{indent} {marker}\\tab {inner}\\par ")
            return "".join(parts) + "\\li0 "
        inner = "".join(render(c) for c in elem.children)
        if tag in ("strong", "b"):
            return f"\\b {inner}\\b0 "
        if tag in ("em", "i"):
            return f"\\i {inner}\\i0 "
        if tag == "u":
            return f"\\ul {inner}\\ulnone "
        if tag in ("p", "div"):
            return f"{inner}\\par "
        return inner

    return "".join(render(c) for c in soup.children)


def render_chapter_rtf(content_html: str) -> bytes:
    body = html_to_rtf_body(content_html)
    return ("{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0 Calibri;}}\\f0\\fs24 " + body + "}").encode("utf-8")


def render_folder_rtf(folder_name: str, chapters: list) -> bytes:
    parts = [f"{{\\b\\fs36 {_rtf_escape(folder_name)}\\par}}\\par "]
    for chapter in chapters:
        parts.append(f"{{\\b\\fs28 {_rtf_escape(chapter.name)}\\par}}")
        parts.append(html_to_rtf_body(chapter.content_html))
        parts.append("\\par ")
    body = "".join(parts)
    return ("{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0 Calibri;}}\\f0\\fs24 " + body + "}").encode("utf-8")


def render_folder_markdown(folder_name: str, chapters: list) -> str:
    parts = [f"# {folder_name}"]
    for chapter in chapters:
        parts.append(f"## {chapter.name}\n\n{html_to_markdown(chapter.content_html)}")
    return "\n\n".join(parts)


def render_folder_txt(folder_name: str, chapters: list) -> str:
    parts = [folder_name]
    for chapter in chapters:
        parts.append(f"{chapter.name}\n\n{html_to_text(chapter.content_html)}")
    return "\n\n\n".join(parts)


def render_folder_docx(folder_name: str, chapters: list) -> Document:
    doc = Document()
    doc.add_heading(folder_name, level=0)
    for idx, chapter in enumerate(chapters):
        doc.add_heading(chapter.name, level=1)
        append_html_to_docx(doc, chapter.content_html)
        if idx < len(chapters) - 1:
            doc.add_page_break()
    return doc


def folder_chapters_recursive(folder_id: uuid.UUID) -> list:
    """All chapters in a folder and its descendants -- both nested
    sub-folders and chapters nested under other chapters -- in a stable
    read order."""
    return (
        Chapter.query.filter(Chapter.id.in_(chapter_ids_under_folder(folder_id)))
        .order_by(Chapter.position, Chapter.created_at)
        .all()
    )


def get_user_settings() -> UserSettings:
    settings = db.session.get(UserSettings, current_user.id)
    if settings is None:
        settings = UserSettings(user_id=current_user.id)
        db.session.add(settings)
        db.session.commit()
    return settings


class HierarchyError(ValueError):
    """One exception type for every hierarchy-rule violation (depth, cycle,
    cross-book move, corrupt stored data) -- every create/move endpoint
    raises the same thing, caught once in api.py and turned into
    error(str(e), 400)."""


MAX_FOLDER_DEPTH = 9   # "~10 folder levels" -- 0-indexed, so depths 0..9
MAX_CHAPTER_DEPTH = 4  # "~5 chapter levels" -- 0-indexed, so depths 0..4
# A walker exceeding this many hops is following corrupt/cyclic stored data
# (a bug, a manual DB edit, anything that bypassed the validators below),
# not a legitimately deep tree -- both caps above are a fraction of this.
HIERARCHY_HOP_CEILING = 64


def descendant_folder_ids(root_folder_id: uuid.UUID) -> list:
    """BFS down Folder.parent_id. Cycle-safe: raises HierarchyError instead
    of looping forever if stored data is ever corrupt -- validate_folder_parent
    is what keeps a *new* cycle from ever being written; this is the
    separate backstop against an existing one taking the app down."""
    ids = [root_folder_id]
    visited = {root_folder_id}
    frontier = [root_folder_id]
    hops = 0
    while frontier:
        hops += 1
        if hops > HIERARCHY_HOP_CEILING:
            raise HierarchyError('Folder hierarchy is corrupt (cycle or excessive depth)')
        children = Folder.query.filter(Folder.parent_id.in_(frontier)).with_entities(Folder.id).all()
        if any(c.id in visited for c in children):
            raise HierarchyError('Folder hierarchy is corrupt (cycle or excessive depth)')
        frontier = [c.id for c in children]
        visited.update(frontier)
        ids.extend(frontier)
    return ids


def descendant_chapter_ids(root_chapter_id: uuid.UUID) -> list:
    """BFS down Chapter.parent_chapter_id -- exact mirror of
    descendant_folder_ids above. Used for: cycle checks, goal-progress
    recursion, stats recursion, search-share expansion, and the cross-book
    book_id cascade on a chapter move."""
    ids = [root_chapter_id]
    visited = {root_chapter_id}
    frontier = [root_chapter_id]
    hops = 0
    while frontier:
        hops += 1
        if hops > HIERARCHY_HOP_CEILING:
            raise HierarchyError('Chapter hierarchy is corrupt (cycle or excessive depth)')
        children = Chapter.query.filter(Chapter.parent_chapter_id.in_(frontier)).with_entities(Chapter.id).all()
        if any(c.id in visited for c in children):
            raise HierarchyError('Chapter hierarchy is corrupt (cycle or excessive depth)')
        frontier = [c.id for c in children]
        visited.update(frontier)
        ids.extend(frontier)
    return ids


def folder_depth(folder: Folder) -> int:
    """0 for a book root, +1 per ancestor folder. Cycle-safe."""
    depth = 0
    visited = {folder.id}
    node = folder
    while node.parent_id is not None:
        if depth > HIERARCHY_HOP_CEILING:
            raise HierarchyError('Folder hierarchy is corrupt (cycle or excessive depth)')
        node = db.session.get(Folder, node.parent_id)
        if node is None or node.id in visited:
            raise HierarchyError('Folder hierarchy is corrupt (cycle or excessive depth)')
        visited.add(node.id)
        depth += 1
    return depth


def chapter_depth(chapter: Chapter) -> int:
    """0 for a chapter directly under a folder, +1 per ancestor chapter."""
    depth = 0
    visited = {chapter.id}
    node = chapter
    while node.parent_chapter_id is not None:
        if depth > HIERARCHY_HOP_CEILING:
            raise HierarchyError('Chapter hierarchy is corrupt (cycle or excessive depth)')
        node = db.session.get(Chapter, node.parent_chapter_id)
        if node is None or node.id in visited:
            raise HierarchyError('Chapter hierarchy is corrupt (cycle or excessive depth)')
        visited.add(node.id)
        depth += 1
    return depth


def folder_subtree_max_depth(folder: Folder) -> int:
    """Depth of the deepest descendant folder relative to `folder` itself
    (0 if it has no child folders)."""
    max_depth = 0
    visited = {folder.id}
    frontier = {folder.id: 0}
    hops = 0
    while frontier:
        hops += 1
        if hops > HIERARCHY_HOP_CEILING:
            raise HierarchyError('Folder hierarchy is corrupt (cycle or excessive depth)')
        children = Folder.query.filter(Folder.parent_id.in_(frontier.keys())).with_entities(
            Folder.id, Folder.parent_id
        ).all()
        next_frontier = {}
        for c in children:
            if c.id in visited:
                raise HierarchyError('Folder hierarchy is corrupt (cycle or excessive depth)')
            visited.add(c.id)
            d = frontier[c.parent_id] + 1
            max_depth = max(max_depth, d)
            next_frontier[c.id] = d
        frontier = next_frontier
    return max_depth


def chapter_subtree_max_depth(chapter: Chapter) -> int:
    """Depth of the deepest descendant chapter relative to `chapter` itself
    (0 if it has no child chapters)."""
    max_depth = 0
    visited = {chapter.id}
    frontier = {chapter.id: 0}
    hops = 0
    while frontier:
        hops += 1
        if hops > HIERARCHY_HOP_CEILING:
            raise HierarchyError('Chapter hierarchy is corrupt (cycle or excessive depth)')
        children = Chapter.query.filter(Chapter.parent_chapter_id.in_(frontier.keys())).with_entities(
            Chapter.id, Chapter.parent_chapter_id
        ).all()
        next_frontier = {}
        for c in children:
            if c.id in visited:
                raise HierarchyError('Chapter hierarchy is corrupt (cycle or excessive depth)')
            visited.add(c.id)
            d = frontier[c.parent_chapter_id] + 1
            max_depth = max(max_depth, d)
            next_frontier[c.id] = d
        frontier = next_frontier
    return max_depth


def nearest_folder_for_chapter(chapter: Chapter) -> Folder | None:
    """Walks parent_chapter_id upward until it finds a chapter with
    folder_id set, then returns that folder -- resolves a nested chapter's
    containing Folder for book-color, breadcrumb, and sharing
    ancestor-walk purposes. None only in the pathological case of a
    chapter with no folder anywhere in its ancestry, which
    chk_chapters_one_parent should make impossible in practice."""
    node = chapter
    visited = {chapter.id}
    hops = 0
    while node.folder_id is None:
        hops += 1
        if hops > HIERARCHY_HOP_CEILING or node.parent_chapter_id is None:
            return None
        node = db.session.get(Chapter, node.parent_chapter_id)
        if node is None or node.id in visited:
            raise HierarchyError('Chapter hierarchy is corrupt (cycle or excessive depth)')
        visited.add(node.id)
    return db.session.get(Folder, node.folder_id)


def chapter_ids_under_folder(folder_id: uuid.UUID) -> list:
    """Every chapter administratively under a folder subtree, INCLUDING
    chapters nested under other chapters within it -- the canonical
    replacement for the old `Chapter.folder_id.in_(descendant_folder_ids(...))`
    pattern, which only ever saw direct folder-child chapters."""
    folder_ids = descendant_folder_ids(folder_id)
    direct = Chapter.query.filter(Chapter.folder_id.in_(folder_ids)).with_entities(Chapter.id).all()
    ids = []
    for c in direct:
        ids.extend(descendant_chapter_ids(c.id))
    return ids


def chapter_ids_with_children(chapter_ids) -> set:
    """Which of the given chapter ids have at least one child chapter --
    one batched query, so a tree listing can decide whether each row needs
    an expand arrow without an N+1 existence-check per chapter."""
    chapter_ids = list(chapter_ids)
    if not chapter_ids:
        return set()
    rows = (
        db.session.query(Chapter.parent_chapter_id)
        .filter(Chapter.parent_chapter_id.in_(chapter_ids))
        .distinct()
        .all()
    )
    return {r[0] for r in rows}


def lock_books_for_hierarchy_change(*book_ids: uuid.UUID) -> None:
    """Acquires a pg_advisory_xact_lock per unique book_id, sorted by
    str(book_id) for deadlock avoidance when a cross-book move locks two
    books at once -- released automatically at transaction end (commit or
    rollback). Call before validate_*_parent in the folder-move and
    chapter-move endpoints; not needed for create/rename (a brand-new node
    can't already be part of a cycle, and a rename doesn't touch
    parent_id/parent_chapter_id). Closes a TOCTOU race: without this, two
    concurrent moves can each validate against pre-move state and jointly
    create a cycle that neither would have permitted alone."""
    for book_id in sorted(set(book_ids), key=str):
        db.session.execute(db.text('SELECT pg_advisory_xact_lock(hashtext(:book_id))'), {'book_id': str(book_id)})


def validate_folder_parent(folder_id_being_moved: uuid.UUID | None, new_parent: Folder) -> None:
    """Raises HierarchyError for a create (folder_id_being_moved is None,
    only the depth check applies) or a move: the folder being moved is a
    book root (can never be reparented -- its own book_id points to itself,
    and "book" only means anything as a stable top-level tree), self-parent,
    new_parent inside the moved folder's own subtree (cycle), a different
    book (folder moves stay same-book), or exceeding MAX_FOLDER_DEPTH once
    the moved subtree's own deepest descendant is accounted for. Permission
    checks (require_folder_access) stay the caller's job -- this only
    checks shape."""
    subtree_depth = 0
    if folder_id_being_moved is not None:
        moving = db.session.get(Folder, folder_id_being_moved)
        if moving is None:
            raise HierarchyError('Folder being moved no longer exists')
        if moving.parent_id is None:
            raise HierarchyError("A book's own root folder can't be moved")
        if new_parent.id == moving.id:
            raise HierarchyError("A folder can't be moved into itself")
        if new_parent.id in descendant_folder_ids(moving.id):
            raise HierarchyError("A folder can't be moved into one of its own sub-folders")
        if new_parent.book_id != moving.book_id:
            raise HierarchyError('Folders can only be moved within the same book')
        subtree_depth = folder_subtree_max_depth(moving)
    if folder_depth(new_parent) + 1 + subtree_depth > MAX_FOLDER_DEPTH:
        raise HierarchyError(f'That would nest folders more than {MAX_FOLDER_DEPTH + 1} levels deep')


def validate_chapter_parent(
    chapter_id_being_moved: uuid.UUID | None,
    *,
    new_folder: Folder | None = None,
    new_parent_chapter: Chapter | None = None,
) -> None:
    """Same shape as validate_folder_parent, for chapters -- exactly one of
    new_folder/new_parent_chapter. Cycle/self-parent checks only apply when
    new_parent_chapter is given (landing directly under a folder can never
    cycle). Cross-book is deliberately NOT checked here -- chapter moves
    are allowed across books (existing behavior); the caller cascades
    book_id onto the moved subtree itself after this validates clean. New
    root depth is 0 when landing directly under a folder, else
    chapter_depth(new_parent_chapter) + 1; reject if that plus the moved
    chapter's own deepest-descendant depth exceeds MAX_CHAPTER_DEPTH."""
    if (new_folder is None) == (new_parent_chapter is None):
        raise HierarchyError('A chapter must have exactly one parent (a folder or a chapter)')

    moving = None
    if chapter_id_being_moved is not None:
        moving = db.session.get(Chapter, chapter_id_being_moved)
        if moving is None:
            raise HierarchyError('Chapter being moved no longer exists')

    if new_parent_chapter is not None:
        if moving is not None:
            if new_parent_chapter.id == moving.id:
                raise HierarchyError("A chapter can't be nested inside itself")
            if new_parent_chapter.id in descendant_chapter_ids(moving.id):
                raise HierarchyError("A chapter can't be nested inside one of its own sub-chapters")
        new_root_depth = chapter_depth(new_parent_chapter) + 1
    else:
        new_root_depth = 0

    subtree_depth = chapter_subtree_max_depth(moving) if moving is not None else 0
    if new_root_depth + subtree_depth > MAX_CHAPTER_DEPTH:
        raise HierarchyError(f'That would nest chapters more than {MAX_CHAPTER_DEPTH + 1} levels deep')


def resource_typed_words(
    *,
    folder: Folder | None = None,
    chapter: Chapter | None = None,
    user_id: uuid.UUID,
    start_date: datetime.date,
    end_date: datetime.date | None = None,
) -> int:
    """Sum of ChapterWritingActivity.words_typed for a goal's target resource
    (a single chapter, or a folder's entire descendant subtree), personal to
    `user_id` and bounded to [start_date, end_date] (end_date inclusive;
    None means no upper bound -- used for a period still in progress, where
    "now" is the bound). This is the engine behind word-goal progress: a
    goal's progress is genuinely-typed activity during its period, not net
    resource word-count growth (see Goal's docstring) -- so pasted content
    and a collaborator's own edits on a shared resource never move it."""
    query = db.session.query(db.func.coalesce(db.func.sum(ChapterWritingActivity.words_typed), 0)).filter(
        ChapterWritingActivity.user_id == user_id,
        ChapterWritingActivity.date >= start_date,
    )
    if end_date is not None:
        query = query.filter(ChapterWritingActivity.date <= end_date)
    if chapter is not None:
        # A chapter-targeted word goal includes its own nested chapters'
        # activity too -- this is the goal's *scope*, distinct from the
        # "chapter's own word count" display (Chapter Stats, editor status
        # bar), which deliberately never auto-aggregates. See Goal's
        # docstring.
        chapter_ids = descendant_chapter_ids(chapter.id)
        query = query.filter(ChapterWritingActivity.chapter_id.in_(chapter_ids))
    elif folder is not None:
        chapter_ids = chapter_ids_under_folder(folder.id)
        if not chapter_ids:
            return 0
        query = query.filter(ChapterWritingActivity.chapter_id.in_(chapter_ids))
    else:
        return 0
    return int(query.scalar() or 0)


def writing_activity_totals(
    chapter_ids,
    *,
    user_id: uuid.UUID | None = None,
    start_date: datetime.date | None = None,
    end_date: datetime.date | None = None,
) -> dict:
    """Reusable activity aggregation for an exact chapter scope.

    Folder/book callers pass chapter_ids_under_folder(); chapter callers pass
    one id unless their UI explicitly promises a subtree. `user_id=None` is
    resource-level across collaborators, while passing a user makes the same
    query reusable for P0.8 contributor breakdowns.
    """
    chapter_ids = list(chapter_ids)
    if not chapter_ids:
        return {'wordsTyped': 0, 'wordsPasted': 0, 'wordsDeleted': 0, 'wordsWritten': 0, 'activeSeconds': 0}
    query = db.session.query(
        db.func.coalesce(db.func.sum(ChapterWritingActivity.words_typed), 0),
        db.func.coalesce(db.func.sum(ChapterWritingActivity.words_pasted), 0),
        db.func.coalesce(db.func.sum(ChapterWritingActivity.words_deleted), 0),
        db.func.coalesce(db.func.sum(ChapterWritingActivity.active_seconds), 0),
    ).filter(ChapterWritingActivity.chapter_id.in_(chapter_ids))
    if user_id is not None:
        query = query.filter(ChapterWritingActivity.user_id == user_id)
    if start_date is not None:
        query = query.filter(ChapterWritingActivity.date >= start_date)
    if end_date is not None:
        query = query.filter(ChapterWritingActivity.date <= end_date)
    typed, pasted, deleted, active = query.one()
    typed = int(typed)
    deleted = int(deleted)
    return {
        'wordsTyped': typed,
        'wordsPasted': int(pasted),
        'wordsDeleted': deleted,
        'wordsWritten': words_written_from(typed, deleted),
        'activeSeconds': int(active),
    }


def writing_activity_contributions(chapter_ids) -> list[dict]:
    """Per-user activity totals for an exact resource scope.

    Deliberately joins activity directly to its original User instead of to
    current shares. A removed collaborator must remain credited in an
    authorized viewer's historical stats even though that collaborator can
    no longer open the resource themselves. WPM is calculated independently
    for each row; there is no meaningful resource-level/blended WPM.
    """
    chapter_ids = list(chapter_ids)
    if not chapter_ids:
        return []
    rows = (
        db.session.query(
            User.id,
            User.username,
            db.func.coalesce(db.func.sum(ChapterWritingActivity.words_typed), 0),
            db.func.coalesce(db.func.sum(ChapterWritingActivity.words_pasted), 0),
            db.func.coalesce(db.func.sum(ChapterWritingActivity.words_deleted), 0),
            db.func.coalesce(db.func.sum(ChapterWritingActivity.active_seconds), 0),
        )
        .join(User, User.id == ChapterWritingActivity.user_id)
        .filter(ChapterWritingActivity.chapter_id.in_(chapter_ids))
        .group_by(User.id, User.username)
        .order_by(User.username, User.id)
        .all()
    )
    contributions = []
    for user_id, username, typed, pasted, deleted, active in rows:
        typed = int(typed)
        pasted = int(pasted)
        deleted = int(deleted)
        active = int(active)
        contributions.append({
            'userId': user_id,
            'username': username,
            'wordsTyped': typed,
            'wordsPasted': pasted,
            'wordsDeleted': deleted,
            'wordsWritten': words_written_from(typed, deleted),
            'activeSeconds': active,
            'wpm': calculate_wpm(typed, active),
        })
    return contributions


def resource_completed_chapter_count(
    folder: Folder,
    period_start: datetime.date,
    user=None,
    period_end: datetime.date | None = None,
) -> int:
    """Chapters under `folder` (including nested sub-folders AND chapters
    nested under other chapters within it) marked complete on or after
    `period_start` -- a chapter completed long before the goal's current
    period doesn't count toward it."""
    user = user or current_user
    cutoff, after_end = local_date_bounds_utc(period_start, period_end or user_local_today(user), user)
    return Chapter.query.filter(
        Chapter.id.in_(chapter_ids_under_folder(folder.id)),
        Chapter.completed_at.isnot(None),
        Chapter.completed_at >= cutoff,
        Chapter.completed_at < after_end,
    ).count()


def record_writing_activity(
    chapter: Chapter,
    user_id: uuid.UUID,
    elapsed_seconds: float,
    typed_words_delta: int,
    pasted_words_delta: int,
    deleted_words_delta: int = 0,
    *,
    occurred_at: datetime.datetime | None = None,
) -> None:
    """Accumulate one heartbeat interval's worth of active writing time and
    input-source-classified words into ChapterWritingActivity, for the
    (user, chapter, today, this hour) bucket. `elapsed_seconds` is expected
    to already be 0 unless genuine typing/deleting input happened this
    interval (see api.api_heartbeat_chapter_presence) -- a paste-only
    interval still credits `pasted_words_delta` but never active_seconds,
    which is what keeps WPM's active-time denominator immune to paste.
    `deleted_words_delta` (P0.11) is genuine keyboard/composition words
    removed this interval -- like typed/pasted, it's still credited even on
    an interval with 0 elapsed active time (e.g. hadTypingInput false but a
    dropped/duplicated heartbeat catching up a nonzero cumulative delta).
    Uses a real Postgres upsert (ON CONFLICT DO UPDATE), so incrementing the
    hourly bucket is atomic as well. It deliberately does not commit: the
    caller holds the matching ChapterPresence row lock and commits presence
    counter advancement plus this activity update together."""
    if elapsed_seconds <= 0 and typed_words_delta <= 0 and pasted_words_delta <= 0 and deleted_words_delta <= 0:
        return
    user = db.session.get(User, user_id)
    local_now = user_local_datetime(user, occurred_at)
    stmt = pg_insert(ChapterWritingActivity).values(
        user_id=user_id,
        chapter_id=chapter.id,
        date=local_now.date(),
        hour_of_day=local_now.hour,
        active_seconds=int(elapsed_seconds),
        words_typed=typed_words_delta,
        words_pasted=pasted_words_delta,
        words_deleted=deleted_words_delta,
    )
    stmt = stmt.on_conflict_do_update(
        index_elements=[
            ChapterWritingActivity.user_id,
            ChapterWritingActivity.chapter_id,
            ChapterWritingActivity.date,
            ChapterWritingActivity.hour_of_day,
        ],
        set_={
            "active_seconds": ChapterWritingActivity.active_seconds + stmt.excluded.active_seconds,
            "words_typed": ChapterWritingActivity.words_typed + stmt.excluded.words_typed,
            "words_pasted": ChapterWritingActivity.words_pasted + stmt.excluded.words_pasted,
            "words_deleted": ChapterWritingActivity.words_deleted + stmt.excluded.words_deleted,
        },
    )
    db.session.execute(stmt)


def compute_writing_streak(
    user_id: uuid.UUID,
    *,
    today: datetime.date | None = None,
    chapter_ids=None,
) -> dict:
    """Current and longest consecutive-day streaks of real writing activity
    (any words_typed > 0 that day), personal to `user_id` -- matches how
    Goal is personal, not resource-shared. Pasted words never land in
    words_typed, so a paste-only day (e.g. importing an old manuscript)
    correctly doesn't create a streak day."""
    query = (
        db.session.query(ChapterWritingActivity.date, db.func.sum(ChapterWritingActivity.words_typed))
        .filter(ChapterWritingActivity.user_id == user_id)
        .group_by(ChapterWritingActivity.date)
        .having(db.func.sum(ChapterWritingActivity.words_typed) > 0)
    )
    if chapter_ids is not None:
        chapter_ids = list(chapter_ids)
        if not chapter_ids:
            return {"current": 0, "longest": 0}
        query = query.filter(ChapterWritingActivity.chapter_id.in_(chapter_ids))
    rows = query.all()
    active_days = sorted(d for d, _ in rows)
    if not active_days:
        return {"current": 0, "longest": 0}

    longest = 1
    run = 1
    for prev, cur in zip(active_days, active_days[1:]):
        if (cur - prev).days == 1:
            run += 1
        else:
            longest = max(longest, run)
            run = 1
    longest = max(longest, run)

    today = today or user_local_today(db.session.get(User, user_id))
    current = 0
    if active_days[-1] in (today, today - datetime.timedelta(days=1)):
        current = 1
        for i in range(len(active_days) - 1, 0, -1):
            if (active_days[i] - active_days[i - 1]).days == 1:
                current += 1
            else:
                break
    return {"current": current, "longest": longest}


def chapter_last_activity_date(chapter: Chapter, user=None) -> datetime.date | None:
    """Best-known date of real writing activity on `chapter`, resource-level
    (every collaborator's activity, no user filter -- matches word-count
    scoping elsewhere). Falls back to updated_at's date for a chapter with no
    ChapterWritingActivity rows yet (the table starts empty at deploy, and
    updated_at is bumped by any column change, not just writing, but it's the
    best signal available until real activity accumulates)."""
    latest = (
        db.session.query(db.func.max(ChapterWritingActivity.date))
        .filter(ChapterWritingActivity.chapter_id == chapter.id)
        .scalar()
    )
    if latest is not None:
        return latest
    return user_local_datetime(user, chapter.updated_at).date() if chapter.updated_at else None


def _goal_period_final_current(goal: Goal, period_start: datetime.date, period_end: datetime.date) -> int:
    """Final progress value for one already-elapsed period, for its
    GoalPeriodHistory row. Unlike the live progress shown for the *current*
    period (which has no upper bound -- "now" is the bound), a past
    period's count is capped at its own end so it doesn't pick up activity
    from periods after it. Since progress is a date-range sum rather than a
    captured baseline, this is well-defined for *any* period, including one
    that was never the "current" period when read -- see
    advance_goal_period, which now records every elapsed period, not just
    the first."""
    if goal.goal_type == GoalType.words:
        return resource_typed_words(
            folder=goal.folder, chapter=goal.chapter, user_id=goal.user_id,
            start_date=period_start, end_date=period_end,
        )
    return resource_completed_chapter_count(
        goal.folder, period_start, user=goal.user, period_end=period_end
    )


def advance_date_by_cadence(period_start: datetime.date, cadence: str) -> datetime.date:
    if cadence == 'daily':
        return period_start + datetime.timedelta(days=1)
    if cadence == 'weekly':
        return period_start + datetime.timedelta(days=7)
    if cadence == 'monthly':
        year = period_start.year + period_start.month // 12
        month = period_start.month % 12 + 1
        day = min(period_start.day, calendar.monthrange(year, month)[1])
        return datetime.date(year, month, day)
    raise ValueError(f'Unknown cadence: {cadence}')


def advance_goal_period(goal: Goal, *, today: datetime.date | None = None) -> None:
    """Lazily roll a recurring goal's period forward past any boundaries
    that have already elapsed. Called on every read of a goal (see
    api.api_list_goals) -- idempotent, mutates `goal` in place, caller is
    responsible for committing. A fixed-range goal (cadence is None) never
    rolls. A recurring goal with an end_date stops rolling once that date
    has passed, freezing it at its last active period.

    Progress is now a live date-range sum (services.resource_typed_words /
    resource_completed_chapter_count) rather than a captured baseline, so
    every period this rolls past -- not just the first -- gets an exact
    GoalPeriodHistory row; there's no more "only the first skipped period
    was ever observed" limitation."""
    today = today or user_local_today(goal.user)
    if goal.cadence is not None and (goal.end_date is None or today <= goal.end_date):
        while today >= advance_date_by_cadence(goal.period_start, goal.cadence.value):
            next_period_start = advance_date_by_cadence(goal.period_start, goal.cadence.value)
            period_end = next_period_start - datetime.timedelta(days=1)
            current = _goal_period_final_current(goal, goal.period_start, period_end)
            db.session.add(GoalPeriodHistory(
                goal_id=goal.id,
                period_start=goal.period_start,
                period_end=period_end,
                target=goal.target,
                current=current,
                achieved=current >= goal.target,
            ))
            goal.period_start = next_period_start


def next_folder_position(parent_id) -> int:
    q = db.session.query(db.func.max(Folder.position))
    q = q.filter(Folder.parent_id.is_(None)) if parent_id is None else q.filter(Folder.parent_id == parent_id)
    return (q.scalar() or -1) + 1


def next_chapter_position(folder_id: uuid.UUID | None = None, parent_chapter_id: uuid.UUID | None = None) -> int:
    """Next sibling position under whichever immediate parent is given --
    exactly one of folder_id/parent_chapter_id, mirroring Chapter's own
    exactly-one-parent constraint."""
    assert (folder_id is None) != (parent_chapter_id is None), 'exactly one of folder_id/parent_chapter_id'
    q = db.session.query(db.func.max(Chapter.position))
    q = q.filter(Chapter.folder_id == folder_id) if folder_id is not None else q.filter(
        Chapter.parent_chapter_id == parent_chapter_id
    )
    return (q.scalar() or -1) + 1


def snapshot_chapter_version(chapter: Chapter, *, force: bool = False) -> None:
    """Write a ChapterVersion checkpoint capturing the chapter's CURRENT
    content_html, before the caller overwrites it. Skipped unless `force` or
    the last snapshot for this chapter is older than CHAPTER_VERSION_MIN_INTERVAL,
    so continuous autosave doesn't create a row per keystroke pause. Prunes
    down to the most recent CHAPTER_VERSION_RETENTION rows afterward."""
    if not force:
        latest = (
            ChapterVersion.query.filter_by(chapter_id=chapter.id)
            .order_by(ChapterVersion.created_at.desc())
            .first()
        )
        if latest is not None:
            now = datetime.datetime.now(datetime.timezone.utc)
            latest_created_at = latest.created_at
            if latest_created_at.tzinfo is None:
                latest_created_at = latest_created_at.replace(tzinfo=datetime.timezone.utc)
            if now - latest_created_at < CHAPTER_VERSION_MIN_INTERVAL:
                return
    db.session.add(ChapterVersion(chapter_id=chapter.id, content_html=chapter.content_html))
    chapter.version_count += 1
    db.session.flush()
    stale_ids = [
        v.id
        for v in ChapterVersion.query.filter_by(chapter_id=chapter.id)
        .order_by(ChapterVersion.created_at.desc())
        .offset(CHAPTER_VERSION_RETENTION)
        .all()
    ]
    if stale_ids:
        ChapterVersion.query.filter(ChapterVersion.id.in_(stale_ids)).delete(synchronize_session=False)


def create_book(
    name: str, owner_id: uuid.UUID, description: str = '', author: str = '', color: str = '',
    book_type: str = 'general',
) -> Folder:
    """Insert a new top-level book. book_id is self-referential (NOT NULL FK to
    folders.id), so generate the id client-side and insert id == book_id
    together in one statement -- Postgres checks FK constraints post-statement,
    so a row referencing itself in a single INSERT is valid.

    `book_type` defaults to 'general' -- the safe default for direct/plain
    creation (see api_create_book); the interactive New Book wizard passes
    'novel' explicitly to preserve its existing default. Callers are
    expected to have already validated against BOOK_TYPES."""
    new_id = uuid.uuid4()
    folder = Folder(
        id=new_id,
        book_id=new_id,
        parent_id=None,
        owner_id=owner_id,
        name=name,
        description=description,
        author=author,
        color=color,
        book_type=book_type,
    )
    db.session.add(folder)
    db.session.flush()
    return folder


# P1.1A: stable format IDs only -- never an arbitrary user-supplied string.
# Kept in sync with users.chk_users_journal_date_format /
# chk_users_journal_time_format (see the migration and models.User).
JOURNAL_DATE_FORMATS = (
    'long_month_day_year', 'short_month_day_year', 'day_long_month_year', 'day_short_month_year',
    'us_numeric', 'day_first_numeric', 'iso', 'weekday_long',
)
JOURNAL_TIME_FORMATS = ('12_hour', '24_hour')
DEFAULT_JOURNAL_DATE_FORMAT = 'long_month_day_year'
DEFAULT_JOURNAL_TIME_FORMAT = '12_hour'


def format_journal_date(date: datetime.date, format_id: str) -> str:
    """Deterministic, English-only formatting for a Journal day Chapter's
    *default* name (see journal_write_today) -- presentation only, never
    touches the stored journal_date. `date` is already the resolved local
    Journal day; this never infers or re-derives it. An unrecognized
    format_id (never possible for a value that passed the User check
    constraint, but defensive) falls back to the default rather than
    raising, since a bad id here would otherwise block Write Today
    entirely."""
    if format_id not in JOURNAL_DATE_FORMATS:
        format_id = DEFAULT_JOURNAL_DATE_FORMAT
    month_name = calendar.month_name[date.month]
    month_abbr = calendar.month_abbr[date.month]
    if format_id == 'long_month_day_year':
        return f"{month_name} {date.day}, {date.year}"
    if format_id == 'short_month_day_year':
        return f"{month_abbr} {date.day}, {date.year}"
    if format_id == 'day_long_month_year':
        return f"{date.day} {month_name} {date.year}"
    if format_id == 'day_short_month_year':
        return f"{date.day} {month_abbr} {date.year}"
    if format_id == 'us_numeric':
        return f"{date.month:02d}/{date.day:02d}/{date.year}"
    if format_id == 'day_first_numeric':
        return f"{date.day:02d}/{date.month:02d}/{date.year}"
    if format_id == 'iso':
        return date.isoformat()
    # weekday_long -- date.weekday() (0=Monday) matches calendar.day_name's
    # own indexing directly, no conversion needed.
    return f"{calendar.day_name[date.weekday()]}, {month_name} {date.day}, {date.year}"


def format_journal_time(instant: datetime.datetime, timezone_name: str, format_id: str) -> str:
    """Deterministic time-of-day label for a Write Today timestamp, in
    `timezone_name` (the Book owner's IANA zone -- see journal_write_today),
    never including seconds or the date itself (the Journal Chapter already
    represents the day). Presentation only, computed fresh every call --
    never stored, and never rewrites a previously inserted timestamp."""
    if format_id not in JOURNAL_TIME_FORMATS:
        format_id = DEFAULT_JOURNAL_TIME_FORMAT
    local = instant.astimezone(timezone_info(timezone_name))
    if format_id == '24_hour':
        return f"{local.hour:02d}:{local.minute:02d}"
    hour_12 = local.hour % 12 or 12
    period = 'AM' if local.hour < 12 else 'PM'
    return f"{hour_12}:{local.minute:02d} {period}"


def journal_write_today(book: Folder) -> tuple[Chapter, bool]:
    """Idempotent "Write Today" resolution for a Journal Book (P1.2).

    "Today" is always the Book *owner's* configured local date (see
    user_local_today), never the requesting collaborator's -- a shared
    Journal has exactly one definition of today regardless of which Editor
    clicks. Falls back to UTC the same safe way every other calendar-aware
    endpoint does if the owner's stored timezone is missing/invalid.

    Every step searches the *entire* Book by metadata (journal_date /
    journal_year+journal_month), never by name or current parent -- an
    existing year/month Folder or day Chapter is reused exactly where it
    currently lives, never renamed or moved back, so user reorganization
    (renaming "August" to "Summer", moving a day Chapter elsewhere) stays
    permanently safe. Only ever creates what's missing.

    Locks the Book for the duration (released at transaction end) so two
    concurrent Write Today requests can't each decide independently that no
    year/month/day exists yet and both try to create one -- the partial
    unique indexes on Folder/Chapter remain the final backstop, but this is
    what keeps that backstop from ever actually firing in practice. Does
    NOT commit; the caller commits once, so a failure partway through
    (year created, month creation fails) leaves nothing behind rather than
    a half-created hierarchy.

    Returns (chapter, created) -- created is False when today's Chapter
    already existed and nothing was made."""
    lock_books_for_hierarchy_change(book.id)
    owner = db.session.get(User, book.owner_id)
    today = user_local_today(owner)

    existing = Chapter.query.filter_by(book_id=book.id, journal_date=today).first()
    if existing is not None:
        return existing, False

    year_folder = Folder.query.filter_by(
        book_id=book.id, journal_year=today.year, journal_month=None,
    ).first()
    if year_folder is None:
        year_folder = Folder(
            parent_id=book.id, book_id=book.id, name=str(today.year),
            journal_year=today.year, journal_month=None,
            position=next_folder_position(book.id),
        )
        db.session.add(year_folder)
        db.session.flush()

    month_folder = Folder.query.filter_by(
        book_id=book.id, journal_year=today.year, journal_month=today.month,
    ).first()
    if month_folder is None:
        month_folder = Folder(
            parent_id=year_folder.id, book_id=book.id, name=calendar.month_name[today.month],
            journal_year=today.year, journal_month=today.month,
            position=next_folder_position(year_folder.id),
        )
        db.session.add(month_folder)
        db.session.flush()

    day_chapter = Chapter(
        folder_id=month_folder.id, book_id=book.id,
        # P1.1A: the owner's *current* date-format preference -- applied
        # only at creation time, as an initial display name (see
        # format_journal_date). A later preference change never renames
        # this or any other already-existing day Chapter.
        name=format_journal_date(today, owner.journal_date_format if owner else DEFAULT_JOURNAL_DATE_FORMAT),
        journal_date=today,
        position=next_chapter_position(folder_id=month_folder.id),
    )
    db.session.add(day_chapter)
    db.session.flush()
    return day_chapter, True


def ordered_accessible_books() -> list:
    ids = accessible_book_ids()
    if not ids:
        return []
    books = Folder.query.filter(Folder.id.in_(ids)).all()
    order = get_user_settings().book_order
    by_id = {b.id: b for b in books}
    ordered = [by_id[i] for i in order if i in by_id]
    seen = set(order)
    remaining = sorted((b for b in books if b.id not in seen), key=lambda b: b.created_at)
    return ordered + remaining


def serialize_chapter(chapter: Chapter) -> dict:
    children = (
        Chapter.query.filter_by(parent_chapter_id=chapter.id).order_by(Chapter.position, Chapter.created_at).all()
    )
    return {
        'source_id': str(chapter.id),
        'name': chapter.name,
        'description': chapter.description,
        'content_html': chapter.content_html,
        'notes_text': chapter.notes_text,
        'completed_at': chapter.completed_at.isoformat() if chapter.completed_at else None,
        'show_book_color': chapter.show_book_color,
        # P1.2 (v4): the Journal day this Chapter represents, if any --
        # authoritative identity independent of name/position, see
        # Chapter.journal_date's docstring.
        'journal_date': chapter.journal_date.isoformat() if chapter.journal_date else None,
        'children': [serialize_chapter(c) for c in children],
    }


def serialize_folder(folder: Folder) -> dict:
    subfolders = Folder.query.filter_by(parent_id=folder.id).order_by(Folder.position, Folder.created_at).all()
    chapters = Chapter.query.filter_by(folder_id=folder.id).order_by(Chapter.position, Chapter.created_at).all()
    return {
        'source_id': str(folder.id),
        'name': folder.name,
        'description': folder.description,
        'author': folder.author,
        'color': folder.color,
        'show_book_color': folder.show_book_color,
        # P1.2 (v4): book_type is only ever non-null on the Book/root row
        # this method is called with at the top of the tree -- every nested
        # sub-folder's own book_type is always None, so this needs no
        # depth-aware branching. journal_year/journal_month are likewise
        # None on an ordinary Folder.
        'book_type': folder.book_type,
        'journal_year': folder.journal_year,
        'journal_month': folder.journal_month,
        'folders': [serialize_folder(s) for s in subfolders],
        'chapters': [serialize_chapter(c) for c in chapters],
    }


def export_books_zip() -> BytesIO:
    """Portable book archive for the current user's accessible library.

    Version 4 preserves book/folder/chapter structure, user-authored
    document metadata, and P1.2 Book Type/Journal metadata, including
    stable source ids used solely to remap internal references on import.
    Accounts, shares, goals, version history, presence, and writing
    telemetry deliberately remain outside this application-level
    portability format; full-state recovery uses pg_dump.
    """
    books = ordered_accessible_books()
    payload = {
        'version': '4.0',
        'exported_at': datetime.datetime.utcnow().isoformat(),
        'exported_by': current_user.username,
        'books': [serialize_folder(b) for b in books],
    }
    mem = BytesIO()
    with zipfile.ZipFile(mem, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('data.json', json.dumps(payload))
    mem.seek(0)
    return mem


def deserialize_chapter(
    data: dict,
    *,
    book_id: uuid.UUID,
    position: int,
    id_map: dict[tuple[str, str], uuid.UUID],
    imported_chapters: list[Chapter],
    folder: Folder | None = None,
    parent_chapter: Chapter | None = None,
) -> Chapter:
    assert (folder is None) != (parent_chapter is None), 'exactly one of folder/parent_chapter'
    chapter = Chapter(
        folder_id=folder.id if folder is not None else None,
        parent_chapter_id=parent_chapter.id if parent_chapter is not None else None,
        book_id=book_id,
        name=clean_name(data.get('name', 'Untitled')) or 'Untitled',
        description=data.get('description', ''),
        content_html=sanitize_html(data.get('content_html', '')),
        notes_text=data.get('notes_text', ''),
        position=position,
        completed_at=_archive_datetime(data.get('completed_at')),
        show_book_color=data.get('show_book_color', True),
        journal_date=_archive_journal_date(data.get('journal_date')),
    )
    db.session.add(chapter)
    db.session.flush()
    source_id = data.get('source_id')
    if source_id:
        id_map[('chapter', source_id)] = chapter.id
    imported_chapters.append(chapter)
    for idx, child in enumerate(data.get('children', [])):
        deserialize_chapter(
            child,
            book_id=book_id,
            position=idx,
            id_map=id_map,
            imported_chapters=imported_chapters,
            parent_chapter=chapter,
        )
    return chapter


def deserialize_folder(
    data: dict,
    parent: Folder,
    book_id: uuid.UUID,
    position: int,
    *,
    id_map: dict[tuple[str, str], uuid.UUID],
    imported_chapters: list[Chapter],
) -> Folder:
    year = _archive_journal_year(data.get('journal_year'))
    month = _archive_journal_month(data.get('journal_month'), year)
    folder = Folder(
        parent_id=parent.id,
        book_id=book_id,
        name=clean_name(data.get('name', 'Untitled')) or 'Untitled',
        description=data.get('description', ''),
        author=data.get('author', ''),
        color=data.get('color', ''),
        position=position,
        show_book_color=data.get('show_book_color', True),
        journal_year=year,
        journal_month=month,
    )
    db.session.add(folder)
    db.session.flush()
    source_id = data.get('source_id')
    if source_id:
        id_map[('folder', source_id)] = folder.id
    for idx, sub in enumerate(data.get('folders', [])):
        deserialize_folder(
            sub, folder, book_id, idx, id_map=id_map, imported_chapters=imported_chapters
        )
    for idx, chap in enumerate(data.get('chapters', [])):
        deserialize_chapter(
            chap,
            book_id=book_id,
            position=idx,
            id_map=id_map,
            imported_chapters=imported_chapters,
            folder=folder,
        )
    return folder


def _archive_datetime(value) -> datetime.datetime | None:
    if value is None:
        return None
    try:
        parsed = datetime.datetime.fromisoformat(value)
    except (TypeError, ValueError):
        raise ValueError('Archive contains an invalid completion timestamp')
    return parsed if parsed.tzinfo is not None else parsed.replace(tzinfo=datetime.timezone.utc)


def _remap_internal_references(html: str, id_map: dict[tuple[str, str], uuid.UUID]) -> str:
    """Rewrite references whose targets were included in this import.

    References to deleted or non-exported resources intentionally retain
    their old identity and therefore stay unavailable, matching their source
    behavior. Visible anchor text is never rewritten.
    """
    soup = BeautifulSoup(html, 'html.parser')
    changed = False
    for anchor in soup.select('a[data-calwriter-target-type][data-calwriter-target-id]'):
        target_type = anchor.get('data-calwriter-target-type')
        source_id = anchor.get('data-calwriter-target-id')
        target_id = id_map.get((target_type, source_id))
        if target_id is None:
            continue
        target_id_text = str(target_id)
        anchor['data-calwriter-target-id'] = target_id_text
        anchor['href'] = f'calwriter://{target_type}/{target_id_text}'
        changed = True
    return sanitize_html(str(soup)) if changed else html


def _archive_journal_date(value) -> datetime.date | None:
    if value is None:
        return None
    try:
        return datetime.date.fromisoformat(value)
    except (TypeError, ValueError):
        raise ValueError('Archive contains an invalid journal date')


def _archive_journal_year(value) -> int | None:
    if value is None:
        return None
    if isinstance(value, bool) or not isinstance(value, int):
        raise ValueError('Archive contains an invalid journal year')
    return value


def _archive_journal_month(value, year: int | None) -> int | None:
    if value is None:
        return None
    if isinstance(value, bool) or not isinstance(value, int) or not (1 <= value <= 12):
        raise ValueError('Archive contains an invalid journal month')
    if year is None:
        raise ValueError('Archive contains a journal month without a journal year')
    return value


def _archive_book_type(value) -> str:
    """Missing/older-archive (pre-v4) book_type defaults to 'general' --
    never inferred from the book's name."""
    if value is None:
        return 'general'
    if not isinstance(value, str) or value not in BOOK_TYPES:
        raise ValueError('Archive contains an invalid book type')
    return value


def _validate_import_chapter_depth(data: dict, *, depth: int, seen_dates: set) -> None:
    if not isinstance(data, dict):
        raise ValueError('Archive contains an invalid chapter')
    if depth > MAX_CHAPTER_DEPTH:
        raise HierarchyError(f'Import contains chapters nested more than {MAX_CHAPTER_DEPTH + 1} levels deep')
    _validate_archive_resource(data, child_keys=('children',))
    journal_date = _archive_journal_date(data.get('journal_date'))
    if journal_date is not None:
        # Mirrors uq_chapters_journal_date -- caught here (before any row
        # exists) rather than as a bare IntegrityError partway through an
        # otherwise-atomic import.
        if journal_date in seen_dates:
            raise ValueError('Archive contains duplicate Journal dates within one book')
        seen_dates.add(journal_date)
    for child in data.get('children', []):
        _validate_import_chapter_depth(child, depth=depth + 1, seen_dates=seen_dates)


def _validate_import_folder_depth(data: dict, *, depth: int, seen_dates: set, seen_periods: set) -> None:
    """Walks an already-parsed export payload's folder/chapter trees and
    raises HierarchyError before any row is created if any book's nesting
    would exceed MAX_FOLDER_DEPTH/MAX_CHAPTER_DEPTH -- import is atomic
    (see import_books_zip: every book is validated before any is created),
    so this has to happen up front rather than failing partway through.
    `seen_dates`/`seen_periods` are fresh per book (see import_books_zip):
    duplicate Journal metadata is only ever a same-book concern, mirroring
    the database invariants' own book_id-scoped uniqueness."""
    if not isinstance(data, dict):
        raise ValueError('Archive contains an invalid folder')
    if depth > MAX_FOLDER_DEPTH:
        raise HierarchyError(f'Import contains folders nested more than {MAX_FOLDER_DEPTH + 1} levels deep')
    _validate_archive_resource(data, child_keys=('folders', 'chapters'))
    year = _archive_journal_year(data.get('journal_year'))
    month = _archive_journal_month(data.get('journal_month'), year)
    if year is not None:
        period = (year, month)
        if period in seen_periods:
            raise ValueError('Archive contains duplicate Journal year/month organization within one book')
        seen_periods.add(period)
    for chap in data.get('chapters', []):
        _validate_import_chapter_depth(chap, depth=0, seen_dates=seen_dates)
    for sub in data.get('folders', []):
        _validate_import_folder_depth(sub, depth=depth + 1, seen_dates=seen_dates, seen_periods=seen_periods)


def _validate_archive_resource(data: dict, *, child_keys: tuple[str, ...]) -> None:
    for key in ('name', 'description', 'author', 'color', 'content_html', 'notes_text'):
        if key in data and not isinstance(data[key], str):
            raise ValueError(f'Archive contains an invalid {key}')
    for key in ('show_book_color',):
        if key in data and not isinstance(data[key], bool):
            raise ValueError(f'Archive contains an invalid {key}')
    if data.get('source_id') is not None:
        try:
            uuid.UUID(data['source_id'])
        except (AttributeError, TypeError, ValueError):
            raise ValueError('Archive contains an invalid source id')
    if data.get('completed_at') is not None:
        _archive_datetime(data['completed_at'])
    for key in child_keys:
        if not isinstance(data.get(key, []), list):
            raise ValueError(f'Archive contains an invalid {key} list')


def import_books_zip(file_storage, owner_id: uuid.UUID) -> int:
    """Import a .calwdb archive. Always creates new books owned by owner_id --
    never merges/overwrites existing books by name. Returns count imported.
    Raises ValueError on invalid archive contents (caller decides how to report)."""
    try:
        with zipfile.ZipFile(file_storage) as zf:
            if 'data.json' not in zf.namelist():
                raise ValueError('Archive is missing data.json')
            payload = json.loads(zf.read('data.json'))
    except zipfile.BadZipFile:
        raise ValueError('File is not a valid archive')
    except json.JSONDecodeError:
        raise ValueError('Archive contents are invalid')

    if not isinstance(payload, dict) or not isinstance(payload.get('books', []), list):
        raise ValueError('Archive contents are invalid')
    books_data = payload.get('books', [])
    # Validate every book's nesting depth, Book Type, and Journal metadata
    # (including within-book duplicate journal dates/periods) before
    # creating anything -- an import is all-or-nothing, matching "failed
    # moves are atomic" applied to import (see _validate_import_folder_depth).
    # Fresh seen_dates/seen_periods per book: duplicates are only ever a
    # same-book concern (mirrors the book_id-scoped database invariants).
    for book_data in books_data:
        _archive_book_type(book_data.get('book_type'))
        _validate_import_folder_depth(book_data, depth=0, seen_dates=set(), seen_periods=set())
    # Keep first-use settings creation in the same transaction as the book
    # rows. get_user_settings() commits a newly-created row immediately,
    # which would otherwise leave that row behind after a later import
    # failure and violate the archive's all-or-nothing contract.
    settings = db.session.get(UserSettings, current_user.id)
    if settings is None:
        settings = UserSettings(user_id=current_user.id)
        db.session.add(settings)
    id_map: dict[tuple[str, str], uuid.UUID] = {}
    imported_chapters: list[Chapter] = []
    try:
        for book_data in books_data:
            name = clean_name(book_data.get('name', 'Untitled')) or 'Untitled'
            base_name, suffix = name, 1
            while Folder.query.filter_by(parent_id=None, name=name).first():
                suffix += 1
                name = f"{base_name} ({suffix})"
            book = create_book(
                name,
                owner_id=owner_id,
                description=book_data.get('description', ''),
                author=book_data.get('author', ''),
                color=book_data.get('color', ''),
                book_type=_archive_book_type(book_data.get('book_type')),
            )
            book.show_book_color = book_data.get('show_book_color', True)
            source_id = book_data.get('source_id')
            if source_id:
                id_map[('book', source_id)] = book.id
            for idx, sub in enumerate(book_data.get('folders', [])):
                deserialize_folder(
                    sub, book, book.id, idx, id_map=id_map, imported_chapters=imported_chapters
                )
            for idx, chap in enumerate(book_data.get('chapters', [])):
                deserialize_chapter(
                    chap,
                    book_id=book.id,
                    position=idx,
                    id_map=id_map,
                    imported_chapters=imported_chapters,
                    folder=book,
                )
            if book.id not in settings.open_book_ids:
                settings.open_book_ids = settings.open_book_ids + [book.id]
        for chapter in imported_chapters:
            chapter.content_html = _remap_internal_references(chapter.content_html, id_map)
        db.session.commit()
    except Exception:
        db.session.rollback()
        raise
    return len(books_data)
