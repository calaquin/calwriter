import os
import datetime
import json
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_from_directory,
    send_file,
    flash,
)
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = 'change-this'

# Application version
VERSION = "0.5.8.1"
app.jinja_env.globals['app_version'] = VERSION

DATA_DIR = os.environ.get('DATA_DIR', os.path.join(os.getcwd(), 'data'))

# Ensure data directory exists
os.makedirs(DATA_DIR, exist_ok=True)

SETTINGS_FILE = os.path.join(DATA_DIR, 'settings.json')
OPEN_BOOKS_FILE = os.path.join(DATA_DIR, 'open_books.json')


def load_open_books():
    if os.path.isfile(OPEN_BOOKS_FILE):
        try:
            with open(OPEN_BOOKS_FILE) as f:
                return json.load(f)
        except json.JSONDecodeError:
            pass
    books = list_all_books()
    save_open_books(books)
    return books


def save_open_books(books: list) -> None:
    with open(OPEN_BOOKS_FILE, 'w') as f:
        json.dump(books, f)


def load_settings():
    if os.path.isfile(SETTINGS_FILE):
        with open(SETTINGS_FILE) as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                pass
    return {
        'dark_mode': False,
        'sidebar_color': '#f0f0f0',
        'text_color': '#000000',
        'bg_color': '#ffffff',
    }


def save_settings(data: dict) -> None:
    with open(SETTINGS_FILE, 'w') as f:
        json.dump(data, f)


def safe_name(name: str) -> str:
    return ''.join(c for c in name if c.isalnum() or c in (' ', '_', '-')).rstrip()


def html_to_text(html: str) -> str:
    """Convert HTML to plain text."""
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator="\n")


def html_to_docx(html: str, path: str) -> None:
    """Save limited HTML content to a DOCX file."""
    doc = Document()
    soup = BeautifulSoup(html, "html.parser")

    def process(elem, paragraph, formatting=None):
        if formatting is None:
            formatting = {}
        if isinstance(elem, str):
            run = paragraph.add_run(elem)
            run.bold = formatting.get("bold", False)
            run.italic = formatting.get("italic", False)
            run.underline = formatting.get("underline", False)
            return
        tag = elem.name
        fmt = formatting.copy()
        if tag in ("strong", "b"):
            fmt["bold"] = True
        if tag in ("em", "i"):
            fmt["italic"] = True
        if tag == "u":
            fmt["underline"] = True
        if tag in ("p", "div", "br"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            for child in elem.children:
                process(child, p, fmt)
            return
        for child in elem.children:
            process(child, paragraph, fmt)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    for child in soup.children:
        process(child, p)
    doc.save(path)


def append_html_to_docx(doc: Document, html: str) -> None:
    """Append HTML content to an existing DOCX document."""
    soup = BeautifulSoup(html, "html.parser")

    def process(elem, paragraph, formatting=None):
        if formatting is None:
            formatting = {}
        if isinstance(elem, str):
            run = paragraph.add_run(elem)
            run.bold = formatting.get("bold", False)
            run.italic = formatting.get("italic", False)
            run.underline = formatting.get("underline", False)
            return
        tag = elem.name
        fmt = formatting.copy()
        if tag in ("strong", "b"):
            fmt["bold"] = True
        if tag in ("em", "i"):
            fmt["italic"] = True
        if tag == "u":
            fmt["underline"] = True
        if tag in ("p", "div", "br"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            for child in elem.children:
                process(child, p, fmt)
            return
        for child in elem.children:
            process(child, paragraph, fmt)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    for child in soup.children:
        process(child, p)


def sanitize_path(folder: str) -> str:
    parts = [safe_name(p) for p in folder.split('/') if p]
    return os.path.join(*parts) if parts else ''


def load_order(folder: str) -> dict:
    """Load ordering info for a folder."""
    order_file = os.path.join(DATA_DIR, sanitize_path(folder), 'order.json')
    if os.path.isfile(order_file):
        try:
            with open(order_file) as f:
                data = json.load(f)
                if 'folders' not in data:
                    data['folders'] = []
                if 'chapters' not in data:
                    data['chapters'] = []
                return data
        except json.JSONDecodeError:
            pass
    return {'folders': [], 'chapters': []}


def save_order(folder: str, order: dict) -> None:
    os.makedirs(os.path.join(DATA_DIR, sanitize_path(folder)), exist_ok=True)
    order_file = os.path.join(DATA_DIR, sanitize_path(folder), 'order.json')
    with open(order_file, 'w') as f:
        json.dump(order, f)


def list_chapters(folder: str):
    path = os.path.join(DATA_DIR, sanitize_path(folder))
    if not os.path.isdir(path):
        return []
    chapters = [
        c
        for c in os.listdir(path)
        if os.path.isdir(os.path.join(path, c))
        and os.path.isfile(os.path.join(path, c, 'chapter.html'))
    ]
    order = load_order(folder).get('chapters', [])
    ordered = [c for c in order if c in chapters]
    remaining = [c for c in chapters if c not in ordered]
    remaining.sort(key=lambda n: os.path.getctime(os.path.join(path, n)))
    return ordered + remaining



def list_subfolders(folder: str):
    path = os.path.join(DATA_DIR, sanitize_path(folder))
    if not os.path.isdir(path):
        return []
    subs = [
        c
        for c in os.listdir(path)
        if os.path.isdir(os.path.join(path, c))
        and not os.path.isfile(os.path.join(path, c, 'chapter.html'))
    ]
    order = load_order(folder).get('folders', [])
    ordered = [s for s in order if s in subs]
    remaining = [s for s in subs if s not in ordered]
    remaining.sort(key=lambda n: os.path.getctime(os.path.join(path, n)))
    return ordered + remaining


def list_all_books():
    path = DATA_DIR
    books = [f for f in os.listdir(path) if os.path.isdir(os.path.join(path, f))]
    order = load_order('').get('folders', [])
    ordered = [b for b in order if b in books]
    remaining = [b for b in books if b not in ordered]
    remaining.sort(key=lambda n: os.path.getctime(os.path.join(path, n)))
    return ordered + remaining


def list_books():
    all_books = list_all_books()
    open_books = load_open_books()
    return [b for b in all_books if b in open_books]


def note_filename(chapter: str) -> str:
    """Return the standard notes filename for a chapter."""
    return f"{chapter.replace(' ', '_')}_notes.txt"


def list_notes(folder: str, chapter: str):
    """Return the notes filename for the chapter if it exists."""
    path = os.path.join(DATA_DIR, sanitize_path(folder), safe_name(chapter))
    filename = note_filename(chapter)
    note_path = os.path.join(path, filename)
    if os.path.isfile(note_path):
        return [filename]
    return []


app.jinja_env.globals['list_chapters'] = list_chapters
app.jinja_env.globals['list_notes'] = list_notes
app.jinja_env.globals['list_subfolders'] = list_subfolders
app.jinja_env.globals['list_books'] = list_books
app.jinja_env.globals['list_all_books'] = list_all_books


@app.context_processor
def inject_app_settings():
    return {'app_settings': load_settings()}


def read_description(folder: str) -> str:
    """Return description text for a folder if present."""
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'description.txt')
    if os.path.isfile(path):
        with open(path) as f:
            return f.read()
    return ''


def write_description(folder: str, text: str) -> None:
    os.makedirs(os.path.join(DATA_DIR, sanitize_path(folder)), exist_ok=True)
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'description.txt')
    with open(path, 'w') as f:
        f.write(text)

def read_author(folder: str) -> str:
    """Return author text for a folder if present."""
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'author.txt')
    if os.path.isfile(path):
        with open(path) as f:
            return f.read()
    return ''

def write_author(folder: str, text: str) -> None:
    os.makedirs(os.path.join(DATA_DIR, sanitize_path(folder)), exist_ok=True)
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'author.txt')
    with open(path, 'w') as f:
        f.write(text)


@app.route('/')
def index():
    all_books = list_all_books()
    open_books = load_open_books()
    folders = [b for b in all_books if b in open_books]
    return render_template(
        'index.html',
        folders=folders,
        all_books=all_books,
        open_books=open_books,
    )


@app.route('/settings', methods=['GET', 'POST'])
def app_settings_page():
    settings = load_settings()
    if request.method == 'POST':
        if 'reset' in request.form:
            settings = {
                'dark_mode': False,
                'sidebar_color': '#f0f0f0',
                'text_color': '#000000',
                'bg_color': '#ffffff',
            }
        else:
            settings['dark_mode'] = bool(request.form.get('dark_mode'))
            settings['sidebar_color'] = request.form.get('sidebar_color', '#f0f0f0') or '#f0f0f0'
            settings['text_color'] = request.form.get('text_color', '#000000') or '#000000'
            settings['bg_color'] = request.form.get('bg_color', '#ffffff') or '#ffffff'
        save_settings(settings)
        flash('Settings saved')
        return redirect(url_for('index'))
    folders = list_books()
    return render_template('app_settings.html', settings=settings, folders=folders)


@app.route('/folder/create', methods=['POST'])
def create_folder():
    name = safe_name(request.form.get('name', ''))
    if not name:
        flash('Folder name required')
        return redirect(url_for('index'))
    path = os.path.join(DATA_DIR, name)
    os.makedirs(path, exist_ok=True)
    order = load_order('')
    if name not in order.get('folders', []):
        order.setdefault('folders', []).append(name)
        save_order('', order)
    open_books = load_open_books()
    if name not in open_books:
        open_books.append(name)
        save_open_books(open_books)
    return redirect(url_for('view_folder', folder=name))


@app.route('/folder/<path:folder>/delete', methods=['POST'])
def delete_folder(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    parent = os.path.dirname(folder_name)
    if os.path.isdir(path):
        import shutil
        shutil.rmtree(path)
        flash('Book deleted')
        if parent:
            order = load_order(parent)
            if os.path.basename(folder_name) in order.get('folders', []):
                order['folders'].remove(os.path.basename(folder_name))
                save_order(parent, order)
    else:
        flash('Book not found')
    if parent:
        return redirect(url_for('view_folder', folder=parent))
    else:
        order = load_order('')
        bname = os.path.basename(folder_name)
        if bname in order.get('folders', []):
            order['folders'].remove(bname)
            save_order('', order)
        open_books = load_open_books()
        if bname in open_books:
            open_books.remove(bname)
            save_open_books(open_books)
    return redirect(url_for('index'))


@app.route('/folder/<path:folder>/close', methods=['POST'])
def close_book(folder):
    """Hide a book from the sidebar."""
    folder_name = sanitize_path(folder)
    open_books = load_open_books()
    if folder_name in open_books:
        open_books.remove(folder_name)
        save_open_books(open_books)
    flash('Book closed')
    return redirect(url_for('index'))


@app.route('/folder/<path:folder>/open', methods=['POST'])
def open_book(folder):
    """Show a previously closed book in the sidebar."""
    folder_name = sanitize_path(folder)
    open_books = load_open_books()
    if folder_name not in open_books:
        open_books.append(folder_name)
        save_open_books(open_books)
    flash('Book opened')
    return redirect(url_for('view_folder', folder=folder_name))


@app.route('/books/reorder', methods=['POST'])
def reorder_books():
    """Reorder top level books."""
    order = load_order('')
    if request.is_json:
        items = request.json.get('order', [])
        order['folders'] = items
        save_order('', order)
        return ('', 204)
    name = request.form.get('item_name')
    direction = request.form.get('direction')
    items = order.get('folders', [])
    if name in items:
        idx = items.index(name)
        if direction == 'up' and idx > 0:
            items[idx], items[idx-1] = items[idx-1], items[idx]
        elif direction == 'down' and idx < len(items)-1:
            items[idx], items[idx+1] = items[idx+1], items[idx]
        order['folders'] = items
        save_order('', order)
    return redirect(url_for('index'))


@app.route('/folder/<path:folder>/settings', methods=['GET', 'POST'])
def folder_settings(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    if not os.path.isdir(path):
        flash('Book not found')
        return redirect(url_for('index'))
    description = read_description(folder_name)
    author = read_author(folder_name)
    order = load_order(folder_name)
    if request.method == 'POST':
        if 'item_type' in request.form:
            typ = request.form['item_type']
            name = request.form['item_name']
            direction = request.form['direction']
            items = order.get(f'{typ}s', [])
            if name in items:
                idx = items.index(name)
                if direction == 'up' and idx > 0:
                    items[idx], items[idx-1] = items[idx-1], items[idx]
                elif direction == 'down' and idx < len(items)-1:
                    items[idx], items[idx+1] = items[idx+1], items[idx]
                order[f'{typ}s'] = items
                save_order(folder_name, order)
            return redirect(url_for('folder_settings', folder=folder_name))
        new_name = safe_name(request.form.get('name', folder_name.split('/')[-1]))
        desc = request.form.get('description', '')
        author_text = request.form.get('author', '')
        if new_name and new_name != folder_name.split('/')[-1]:
            new_path = os.path.join(DATA_DIR, os.path.dirname(folder_name), new_name)
            if os.path.exists(new_path):
                flash('Name already exists')
            else:
                os.rename(path, new_path)
                parent = os.path.dirname(folder_name)
                if parent:
                    parent_order = load_order(parent)
                    old = folder_name.split('/')[-1]
                    if old in parent_order.get('folders', []):
                        idx = parent_order['folders'].index(old)
                        parent_order['folders'][idx] = new_name
                        save_order(parent, parent_order)
                else:
                    root_order = load_order('')
                    old = folder_name.split('/')[-1]
                    if old in root_order.get('folders', []):
                        idx = root_order['folders'].index(old)
                        root_order['folders'][idx] = new_name
                        save_order('', root_order)
                    open_books = load_open_books()
                    if old in open_books:
                        open_books[open_books.index(old)] = new_name
                        save_open_books(open_books)
                folder_name = os.path.join(os.path.dirname(folder_name), new_name).strip('/')
                path = new_path
                flash('Book renamed')
        write_description(folder_name, desc)
        write_author(folder_name, author_text)
        return redirect(url_for('view_folder', folder=folder_name))
    subfolders = list_subfolders(folder_name)
    chapters = list_chapters(folder_name)
    folders = list_books()
    return render_template(
        'settings.html',
        folder=folder_name,
        name=folder_name.split('/')[-1],
        description=description,
        author=author,
        subfolders=subfolders,
        chapters=chapters,
        folders=folders,
    )


@app.route('/folder/<path:folder>/reorder', methods=['POST'])
def reorder_folder(folder):
    folder_name = sanitize_path(folder)
    order = load_order(folder_name)
    if request.is_json:
        typ = request.json.get('type')
        items = request.json.get('order', [])
        if typ in ('folder', 'chapter'):
            order[f'{typ}s'] = items
            save_order(folder_name, order)
        return ('', 204)
    return redirect(url_for('folder_settings', folder=folder_name))


@app.route('/folder/<path:folder>')
def view_folder(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    if not os.path.isdir(path):
        flash('Folder not found')
        return redirect(url_for('index'))
    chapters = list_chapters(folder_name)
    subfolders = list_subfolders(folder_name)
    description = read_description(folder_name)
    author = read_author(folder_name)
    folders = list_books()
    return render_template('folder.html', folder=folder_name, chapters=chapters, subfolders=subfolders, folders=folders, description=description, author=author)


@app.route('/folder/<path:folder>/chapter/create', methods=['POST'])
def create_chapter(folder):
    folder_name = sanitize_path(folder)
    chapter = safe_name(request.form.get('name', ''))
    if not chapter:
        flash('Chapter name required')
        return redirect(url_for('view_folder', folder=folder_name))
    path = os.path.join(DATA_DIR, folder_name, chapter)
    os.makedirs(path, exist_ok=True)
    open(os.path.join(path, 'chapter.html'), 'a').close()
    order = load_order(folder_name)
    if chapter not in order.get('chapters', []):
        order.setdefault('chapters', []).append(chapter)
        save_order(folder_name, order)
    return redirect(url_for('view_chapter', folder=folder_name, chapter=chapter))




@app.route('/folder/<path:folder>/chapter/<chapter>')
def view_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    if not os.path.isdir(path):
        flash('Chapter not found')
        return redirect(url_for('view_folder', folder=folder_name))
    chapter_file = os.path.join(path, 'chapter.html')
    chapter_html = ''
    if os.path.isfile(chapter_file):
        with open(chapter_file) as f:
            chapter_html = f.read()

    notes_file = os.path.join(path, note_filename(chapter_name))
    notes_text = ''
    if os.path.isfile(notes_file):
        with open(notes_file) as f:
            notes_text = f.read()

    folders = list_books()
    chapters = list_chapters(folder_name)
    return render_template(
        'chapter.html',
        folder=folder_name,
        chapter=chapter_name,
        notes_text=notes_text,
        folders=folders,
        chapters=chapters,
        chapter_html=chapter_html
    )





@app.route('/folder/<path:folder>/chapter/<chapter>/notes/save', methods=['POST'])
def save_notes(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    text = request.form.get('notes', '')
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    os.makedirs(path, exist_ok=True)
    note_path = os.path.join(path, note_filename(chapter_name))
    with open(note_path, 'w') as f:
        f.write(text)
    return ('', 204)


@app.route('/folder/<path:folder>/chapter/<chapter>/save', methods=['POST'])
def save_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    text = request.form.get('text', '')
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    os.makedirs(path, exist_ok=True)
    html_path = os.path.join(path, 'chapter.html')
    with open(html_path, 'w') as f:
        f.write(text)
    docx_path = os.path.join(path, 'chapter.docx')
    html_to_docx(text, docx_path)
    return redirect(url_for('view_chapter', folder=folder_name, chapter=chapter_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/autosave', methods=['POST'])
def autosave_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    text = request.form.get('text', '')
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    os.makedirs(path, exist_ok=True)
    html_path = os.path.join(path, 'chapter.html')
    with open(html_path, 'w') as f:
        f.write(text)
    docx_path = os.path.join(path, 'chapter.docx')
    html_to_docx(text, docx_path)
    return ('', 204)


@app.route('/folder/<path:folder>/chapter/<chapter>/delete', methods=['POST'])
def delete_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    if os.path.isdir(path):
        import shutil
        shutil.rmtree(path)
        flash('Chapter deleted')
        order = load_order(folder_name)
        if chapter_name in order.get('chapters', []):
            order['chapters'].remove(chapter_name)
            save_order(folder_name, order)
    else:
        flash('Chapter not found')
    return redirect(url_for('view_folder', folder=folder_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/rename', methods=['POST'])
def rename_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    new_name = safe_name(request.form.get('new_name', ''))
    if not new_name:
        flash('New name required')
        return redirect(url_for('folder_settings', folder=folder_name))
    old_path = os.path.join(DATA_DIR, folder_name, chapter_name)
    new_path = os.path.join(DATA_DIR, folder_name, new_name)
    if os.path.exists(new_path):
        flash('Name already exists')
    else:
        os.rename(old_path, new_path)
        order = load_order(folder_name)
        if chapter_name in order.get('chapters', []):
            idx = order['chapters'].index(chapter_name)
            order['chapters'][idx] = new_name
            save_order(folder_name, order)
        flash('Chapter renamed')
    return redirect(url_for('folder_settings', folder=folder_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/notes/download')
def download_note(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    note_name = note_filename(chapter_name)
    return send_from_directory(path, note_name, as_attachment=True, download_name=note_name)


@app.route('/folder/<path:folder>/chapter/<chapter>/chapter.docx')
def download_chapter_docx(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    return send_from_directory(
        path,
        'chapter.docx',
        as_attachment=True,
        download_name=f"{chapter_name}.docx",
    )




@app.route('/folder/<path:folder>/combined.docx')
def download_combined_docx(folder):
    """Download a DOCX containing all chapters in a folder."""
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    if not os.path.isdir(path):
        flash('Folder not found')
        return redirect(url_for('index'))
    chapters = list_chapters(folder_name)
    if not chapters:
        flash('No chapters to combine')
        return redirect(url_for('view_folder', folder=folder_name))
    doc = Document()
    for idx, chap in enumerate(chapters):
        doc.add_heading(chap, level=1)
        html_file = os.path.join(path, chap, 'chapter.html')
        if os.path.isfile(html_file):
            with open(html_file) as f:
                append_html_to_docx(doc, f.read())
        if idx < len(chapters) - 1:
            doc.add_page_break()
    from io import BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"{folder_name.split('/')[-1]}_combined.docx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/folder/<path:folder>/folder/create', methods=['POST'])
def create_subfolder(folder):
    folder_name = sanitize_path(folder)
    name = safe_name(request.form.get('name', ''))
    if not name:
        flash('Folder name required')
        return redirect(url_for('view_folder', folder=folder_name))
    path = os.path.join(DATA_DIR, folder_name, name)
    os.makedirs(path, exist_ok=True)
    order = load_order(folder_name)
    if name not in order.get('folders', []):
        order.setdefault('folders', []).append(name)
        save_order(folder_name, order)
    return redirect(url_for('view_folder', folder=f"{folder_name}/{name}"))


@app.route('/folder/<path:folder>/rename_subfolder/<sub>', methods=['POST'])
def rename_subfolder(folder, sub):
    folder_name = sanitize_path(folder)
    sub_name = safe_name(sub)
    new_name = safe_name(request.form.get('new_name', ''))
    if not new_name:
        flash('New name required')
        return redirect(url_for('folder_settings', folder=folder_name))
    old_path = os.path.join(DATA_DIR, folder_name, sub_name)
    new_path = os.path.join(DATA_DIR, folder_name, new_name)
    if os.path.exists(new_path):
        flash('Name already exists')
    else:
        os.rename(old_path, new_path)
        order = load_order(folder_name)
        if sub_name in order.get('folders', []):
            idx = order['folders'].index(sub_name)
            order['folders'][idx] = new_name
            save_order(folder_name, order)
        flash('Sub-folder renamed')
    return redirect(url_for('folder_settings', folder=folder_name))


@app.route('/folder/<path:folder>/stats')
def folder_stats(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    days = int(request.args.get('days', 7))
    total_words = 0
    words_per_day = {}
    for root, dirs, files in os.walk(path):
        if 'chapter.html' in files:
            html_path = os.path.join(root, 'chapter.html')
            with open(html_path) as f:
                text = html_to_text(f.read())
            count = len(text.split())
            total_words += count
            day = datetime.date.fromtimestamp(os.path.getmtime(html_path)).isoformat()
            words_per_day[day] = words_per_day.get(day, 0) + count
    sorted_days = sorted(words_per_day.keys())
    if days > 0:
        cutoff = (datetime.date.today() - datetime.timedelta(days=days-1)).isoformat()
        sorted_days = [d for d in sorted_days if d >= cutoff]
    chart_labels = sorted(sorted_days)
    chart_data = [words_per_day[d] for d in chart_labels]
    folders = list_books()
    return render_template(
        'stats.html',
        folder=folder_name,
        total_words=total_words,
        words_per_day=words_per_day,
        folders=folders,
        chart_labels=chart_labels,
        chart_data=chart_data,
        days=days,
    )


@app.route('/changelog')
def changelog_page():
    """Display the changelog markdown."""
    path = os.path.join(os.path.dirname(__file__), 'CHANGELOG.md')
    content = ''
    if os.path.isfile(path):
        with open(path) as f:
            content = f.read()
    folders = list_books()
    return render_template('changelog.html', content=content, folders=folders)


@app.route('/search')
def search():
    query = request.args.get('q', '').strip()
    results = []
    if query:
        qlower = query.lower()
        for root, dirs, files in os.walk(DATA_DIR):
            rel = os.path.relpath(root, DATA_DIR)
            if 'chapter.html' in files:
                chap = os.path.basename(root)
                with open(os.path.join(root, 'chapter.html')) as f:
                    text = html_to_text(f.read())
                if qlower in text.lower():
                    results.append({'folder': rel, 'chapter': chap, 'type': 'chapter'})
            for fn in files:
                if fn.endswith('_notes.txt'):
                    chap = os.path.basename(root)
                    with open(os.path.join(root, fn)) as nf:
                        text = nf.read()
                    if qlower in text.lower():
                        results.append({'folder': rel, 'chapter': chap, 'type': 'notes'})
    folders = list_books()
    return render_template('search.html', q=query, results=results, folders=folders)


@app.route('/help')
def help_page():
    """Display a basic help page."""
    folders = list_books()
    return render_template('help.html', folders=folders)


@app.route('/about')
def about_page():
    """Show license information."""
    folders = list_books()
    return render_template('about.html', folders=folders)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
